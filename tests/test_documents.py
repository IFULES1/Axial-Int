"""P2 tests: chunking logic and route wiring (no OpenAI/Qdrant needed)."""
from fastapi.testclient import TestClient

from app.main import app
from app.modules.documents.extract import chunk_text

client = TestClient(app)


def test_chunk_short_text_single_chunk():
    assert chunk_text("hello world") == ["hello world"]


def test_chunk_empty():
    assert chunk_text("") == []
    assert chunk_text("   ") == []


def test_chunk_overlap_and_coverage():
    words = " ".join(f"word{i}" for i in range(600))  # well over CHUNK_SIZE chars
    chunks = chunk_text(words, size=1000, overlap=200)
    assert len(chunks) > 1
    # Every chunk is within bounds and non-empty.
    assert all(0 < len(c) <= 1000 for c in chunks)
    # Overlap: consecutive chunks share some trailing/leading content.
    joined = " ".join(chunks)
    assert "word0" in joined and "word599" in joined


def test_documents_routes_mounted():
    paths = client.get("/openapi.json").json()["paths"]
    assert "/documents/upload" in paths
    assert "/documents" in paths
    assert "/documents/{doc_id}" in paths
    assert "/rag/search" in paths
    # Protected: no token → 401/403.
    assert client.get("/documents").status_code in (401, 403)
    assert client.post("/rag/search", json={"query": "x"}).status_code in (401, 403)


# --- Robustesse ingestion : formats, tailles, erreurs claires -----------------

def test_extract_docx_paragraphs_and_tables():
    import io as _io

    from docx import Document as DocxDocument

    from app.modules.documents.extract import extract_text

    d = DocxDocument()
    d.add_paragraph("Pitch deck Lumière Quantum")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "ARR"
    t.rows[0].cells[1].text = "1,2 M€"
    buf = _io.BytesIO()
    d.save(buf)
    text = extract_text("deck.docx", buf.getvalue())
    assert "Lumière Quantum" in text
    assert "ARR | 1,2 M€" in text


def test_extract_xlsx_rows():
    import io as _io

    from openpyxl import Workbook

    from app.modules.documents.extract import extract_text

    wb = Workbook()
    ws = wb.active
    ws.title = "Prévisions"
    ws.append(["Mois", "MRR"])
    ws.append(["Janvier", 12000])
    buf = _io.BytesIO()
    wb.save(buf)
    text = extract_text("previsions.xlsx", buf.getvalue())
    assert "Feuille : Prévisions" in text
    assert "Janvier | 12000" in text


def test_ingest_rejects_unsupported_format_and_size(db=None):
    import pytest

    from app.errors import AppError
    from app.modules.documents import service

    with pytest.raises(AppError) as e1:
        service.ingest(None, user_id="u", filename="virus.exe", data=b"x")
    assert e1.value.code == "unsupported_format"

    with pytest.raises(AppError) as e2:
        service.ingest(None, user_id="u", filename="gros.pdf",
                       data=b"0" * (service.MAX_UPLOAD_BYTES + 1))
    assert e2.value.code == "file_too_large"

    with pytest.raises(AppError) as e3:
        service.ingest(None, user_id="u", filename="vide.txt", data=b"")
    assert e3.value.code == "empty_file"


def test_scanned_pdf_gets_a_clear_message(monkeypatch):
    import pytest

    from app.errors import AppError
    from app.modules.documents import service

    # Simule un PDF scanné : extraction (texte + OCR) qui ne trouve rien.
    monkeypatch.setattr("app.modules.documents.service.extract_text", lambda f, d: "")
    with pytest.raises(AppError) as e:
        service.ingest(None, user_id="u", filename="scan.pdf", data=b"%PDF-1.4 fake")
    assert e.value.code == "extraction_failed"
    assert "scanné" in e.value.message
