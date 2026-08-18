"""Text extraction and chunking — pure functions, no I/O beyond parsing bytes.

Chunking mirrors the legacy pipeline: ~1000-char windows with 200-char overlap,
split on whitespace so we don't cut words.
"""
from __future__ import annotations

import io

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


# Extensions acceptées à l'upload. Tout le reste est refusé avec un message clair.
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md")

OCR_MAX_PAGES = 20  # cap OCR cost on big scanned decks


def extract_pdf_text(data: bytes) -> str:
    """Extract text from a PDF byte string. Returns "" if nothing readable."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts).strip()


def ocr_pdf_text(data: bytes) -> str | None:
    """OCR fallback for scanned PDFs (pitch decks are often image-only).

    Returns the recognized text, or None when the OCR stack (pdf2image/poppler +
    pytesseract/tesseract) is not installed — callers turn that into a clear
    user-facing message instead of a crash.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        return None
    try:
        pages = convert_from_bytes(data, dpi=200, last_page=OCR_MAX_PAGES)
        parts = [pytesseract.image_to_string(p, lang="fra+eng") for p in pages]
        return "\n\n".join(t for t in parts if t.strip()).strip()
    except Exception:
        # Poppler/tesseract binary missing, corrupt file… → same graceful signal.
        return None


def extract_docx_text(data: bytes) -> str:
    """Paragraphs + table cells of a .docx, in document order (best effort)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_xlsx_text(data: bytes) -> str:
    """Rows of every sheet as ' | '-joined lines (values only, capped)."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Feuille : {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 2000:  # cap pathological sheets
                break
            cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def _sanitize(text: str) -> str:
    """Strip lone surrogates / un-encodable code points left by bad PDF glyph
    decoding, so downstream JSON + UTF-8 encoding (e.g. to Cohere) never fails."""
    return text.encode("utf-8", "ignore").decode("utf-8")


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file type: PDF (with OCR fallback for scans), DOCX, XLSX,
    CSV/TXT/MD as UTF-8 text."""
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        text = extract_pdf_text(data)
        if not text:
            # Scanned PDF: no text layer → OCR. None = OCR stack unavailable,
            # "" = OCR ran but found nothing; both fall through as empty.
            text = ocr_pdf_text(data) or ""
        return _sanitize(text)
    if lowered.endswith(".docx"):
        return _sanitize(extract_docx_text(data))
    if lowered.endswith(".xlsx"):
        return _sanitize(extract_xlsx_text(data))
    # CSV / TXT / MD — decode as UTF-8 text.
    try:
        return _sanitize(data.decode("utf-8", errors="ignore").strip())
    except Exception:
        return ""


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping windows on word boundaries."""
    text = " ".join(text.split())  # normalize whitespace
    if not text:
        return []
    if len(text) <= size:
        return [text]

    chunks: list[str] = []
    start = 0
    step = max(1, size - overlap)
    while start < len(text):
        end = start + size
        window = text[start:end]
        # Avoid cutting mid-word: back off to the last space if we're not at the end.
        if end < len(text):
            last_space = window.rfind(" ")
            if last_space > size // 2:
                window = window[:last_space]
        chunks.append(window.strip())
        start += step
    return [c for c in chunks if c]
